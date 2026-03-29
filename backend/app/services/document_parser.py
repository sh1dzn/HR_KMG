"""
Document parser — extract text from uploaded files (PDF, DOCX, XLSX, TXT, etc.)
"""
import io
import logging

logger = logging.getLogger("hr_ai.document_parser")


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract text content from an uploaded file."""
    name_lower = filename.lower()

    if name_lower.endswith(".pdf"):
        return _parse_pdf(file_bytes)
    elif name_lower.endswith(".docx"):
        return _parse_docx(file_bytes)
    elif name_lower.endswith(".xlsx") or name_lower.endswith(".xls"):
        return _parse_xlsx(file_bytes)
    elif name_lower.endswith(".txt") or name_lower.endswith(".md") or name_lower.endswith(".csv"):
        return file_bytes.decode("utf-8", errors="replace")
    else:
        # Try as plain text
        return file_bytes.decode("utf-8", errors="replace")


def _parse_pdf(data: bytes) -> str:
    from PyPDF2 import PdfReader
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def _parse_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def _parse_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines = []
    for ws in wb.worksheets:
        lines.append(f"## {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)
